#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════╗
║         Document Indexer — PDF/DOCX → Qdrant Vector DB          ║
║         Production-grade pipeline optimized for RAG             ║
╚══════════════════════════════════════════════════════════════════╝

Decisiones técnicas clave (ver README para detalle):
  - PyMuPDF (fitz):      extracción PDF robusta con layout awareness
  - python-docx:         parse nativo de DOCX preservando estructura
  - RecursiveCharacterTextSplitter: chunking semántico por jerarquía
  - BAAI/bge-large-en-v1.5: SOTA embedding para recuperación (RAG)
  - Qdrant + HNSW:       ANN index de alta precisión y bajo latencia
"""

import os
import re
import sys
import json
import shutil
import hashlib
import logging
import unicodedata
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Generator, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http import models as qdrant_models
from tqdm import tqdm


# ══════════════════════════════════════════════════════════════════
# CONFIGURACIÓN CENTRAL
# Todos los parámetros en un solo lugar para facilitar ajustes
# ══════════════════════════════════════════════════════════════════

@dataclass
class Config:
    """Configuración global del indexador."""

    # ── Rutas ────────────────────────────────────────────────────
    input_folder:    Path = field(default_factory=lambda: Path("Documentos"))
    error_folder:    Path = field(default_factory=lambda: Path("errores"))
    log_file:        Path = field(default_factory=lambda: Path("indexer.log"))
    state_file:      Path = field(default_factory=lambda: Path(".indexer_state.json"))

    # ── Qdrant ───────────────────────────────────────────────────
    qdrant_host:       str = "localhost"
    qdrant_port:       int = 6333
    collection_name:   str = "documents"

    # ── Modelo de embeddings ──────────────────────────────────────
    # BAAI/bge-large-en-v1.5 → 768 dims, #1 en MTEB retrieval (inglés)
    # Alternativa multilingüe: intfloat/multilingual-e5-large (1024 dims)
    embedding_model: str = "BAAI/bge-large-en-v1.5"
    embedding_dim:   int = 1024
    batch_size:      int = 32   # Ajustar según VRAM disponible

    # ── Chunking ─────────────────────────────────────────────────
    # 512 chars (~128 tokens) → sweet spot recall/precision en RAG
    # Overlap 10% → preserva contexto en fronteras sin duplicar demasiado
    chunk_size:       int = 512
    chunk_overlap:    int = 52
    min_chunk_length: int = 60   # Chunks menores se descartan (ruido)

    # ── HNSW (Qdrant) ────────────────────────────────────────────
    # m=16 → balance memoria/recall. Subir a 32 para máxima calidad.
    # ef_construct=200 → índice más preciso, más lento en build-time.
    hnsw_m:            int = 16
    hnsw_ef_construct: int = 200

    # ── Procesamiento ─────────────────────────────────────────────
    skip_already_indexed: bool = True   # Evita re-indexar usando hash MD5
    upsert_batch_size:    int = 128     # Puntos por batch hacia Qdrant


# Instancia global de configuración
CONFIG = Config()


# ══════════════════════════════════════════════════════════════════
# LOGGING — doble handler: archivo + consola con formato claro
# ══════════════════════════════════════════════════════════════════

def setup_logging() -> logging.Logger:
    """Configura logging con salida a consola y archivo."""
    fmt_console = "%(asctime)s  %(levelname)-8s  %(message)s"
    fmt_file    = "%(asctime)s | %(levelname)-8s | %(funcName)-25s | %(message)s"
    datefmt     = "%H:%M:%S"

    logger = logging.getLogger("indexer")
    logger.setLevel(logging.DEBUG)

    # Handler consola (INFO+)
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    ch.setFormatter(logging.Formatter(fmt_console, datefmt=datefmt))

    # Handler archivo (DEBUG+ — captura todo para diagnóstico)
    fh = logging.FileHandler(CONFIG.log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(fmt_file, datefmt="%Y-%m-%d %H:%M:%S"))

    logger.addHandler(ch)
    logger.addHandler(fh)
    return logger


log = setup_logging()


# ══════════════════════════════════════════════════════════════════
# ESTADO PERSISTENTE — evita re-indexar archivos ya procesados
# ══════════════════════════════════════════════════════════════════

def load_state() -> dict:
    """Carga el estado previo de indexación (hashes procesados)."""
    if CONFIG.state_file.exists():
        try:
            with open(CONFIG.state_file, "r") as f:
                return json.load(f)
        except Exception:
            log.warning("No se pudo leer el estado previo. Empezando desde cero.")
    return {}


def save_state(state: dict) -> None:
    """Persiste el estado de indexación."""
    with open(CONFIG.state_file, "w") as f:
        json.dump(state, f, indent=2)


def file_md5(path: Path) -> str:
    """Calcula el hash MD5 de un archivo para detectar cambios."""
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# ══════════════════════════════════════════════════════════════════
# EXTRACCIÓN DE TEXTO
# ══════════════════════════════════════════════════════════════════

def normalize_text(text: str) -> str:
    """
    Limpieza de texto extraído:
    - Normaliza Unicode (NFC elimina caracteres compuestos duplicados)
    - Elimina caracteres de control excepto saltos de línea
    - Colapsa espacios en blanco excesivos
    - Preserva estructura de párrafos (doble newline)
    """
    # Normalización Unicode NFC
    text = unicodedata.normalize("NFC", text)

    # Eliminar caracteres de control (excepto \n y \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normalizar guiones y comillas tipográficas
    text = text.replace("\u2019", "'").replace("\u2018", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")

    # Colapsar múltiples espacios en uno (preservando newlines)
    text = re.sub(r"[ \t]+", " ", text)

    # Colapsar más de 2 newlines seguidos en párrafo
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Eliminar líneas que solo contienen espacios
    text = re.sub(r"^\s+$", "", text, flags=re.MULTILINE)

    return text.strip()


def extract_pdf(path: Path) -> Generator[dict, None, None]:
    """
    Extrae texto de PDF página a página usando PyMuPDF (fitz).

    Ventajas sobre pdfminer/pypdf2:
    - Maneja PDFs con layouts complejos (columnas, tablas)
    - Mejor gestión de encoding y fuentes embebidas
    - Preserva orden de lectura con sort=True
    - Detecta páginas vacías (scanned sin OCR)

    Yields: dict con texto y metadatos de página
    """
    doc = fitz.open(str(path))
    total_pages = len(doc)

    log.debug(f"  PDF '{path.name}': {total_pages} páginas")

    for page_num, page in enumerate(doc, start=1):
        # get_text("text", sort=True) → respeta orden visual de lectura
        raw_text = page.get_text("text", sort=True)
        text = normalize_text(raw_text)

        if len(text) < CONFIG.min_chunk_length:
            log.debug(f"  Página {page_num} vacía o con muy poco texto, omitida")
            continue

        yield {
            "text": text,
            "page": page_num,
            "total_pages": total_pages,
        }

    doc.close()


def extract_docx(path: Path) -> Generator[dict, None, None]:
    """
    Extrae texto de DOCX preservando estructura semántica.

    Estrategia:
    - Agrupa párrafos en bloques temáticos (preserva contexto)
    - Incluye tablas como texto tabulado
    - Descarta headers/footers repetitivos
    - Respeta heading styles para detección de secciones

    Yields: dict con texto y metadatos de sección
    """
    doc = DocxDocument(str(path))
    current_section: list[str] = []
    current_heading: str = ""
    section_idx: int = 0

    def flush_section():
        """Emite la sección actual si tiene contenido suficiente."""
        nonlocal section_idx
        block = "\n".join(current_section).strip()
        block = normalize_text(block)
        if len(block) >= CONFIG.min_chunk_length:
            section_idx += 1
            return {
                "text": block,
                "section": current_heading or f"Sección {section_idx}",
                "section_idx": section_idx,
            }
        return None

    for element in doc.paragraphs:
        style = element.style.name if element.style else ""
        text  = element.text.strip()

        if not text:
            continue

        # Detectar headings como fronteras de sección semántica
        is_heading = style.lower().startswith("heading") or style.lower().startswith("title")

        if is_heading and current_section:
            result = flush_section()
            if result:
                yield result
            current_section = []
            current_heading = text
        else:
            current_section.append(text)

    # Extraer tablas como texto estructurado
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            current_section.append("\n".join(rows))

    # Emitir última sección
    if current_section:
        result = flush_section()
        if result:
            yield result


def extract_text_blocks(path: Path) -> list[dict]:
    """
    Dispatcher: elige el extractor correcto según extensión.
    Retorna lista de bloques con texto + metadatos.
    """
    ext = path.suffix.lower()
    blocks: list[dict] = []

    if ext == ".pdf":
        for block in extract_pdf(path):
            blocks.append(block)
    elif ext == ".docx":
        for block in extract_docx(path):
            blocks.append(block)
    else:
        raise ValueError(f"Formato no soportado: {ext}")

    if not blocks:
        raise ValueError("El documento no contiene texto extraíble (posiblemente escaneado sin OCR)")

    return blocks


# ══════════════════════════════════════════════════════════════════
# CHUNKING — estrategia semántica jerárquica
# ══════════════════════════════════════════════════════════════════

# Separadores ordenados de mayor a menor semántica:
# párrafo → oración → coma/punto → espacio → carácter
# RecursiveCharacterTextSplitter prueba cada separador
# y sólo baja al siguiente si el chunk sigue siendo demasiado grande.
SEMANTIC_SEPARATORS = [
    "\n\n",    # Párrafos (máxima semántica)
    "\n",      # Saltos de línea
    ". ",      # Fin de oración
    "? ",      # Pregunta
    "! ",      # Exclamación
    "; ",      # Punto y coma
    ", ",      # Coma
    " ",       # Palabras
    "",        # Caracteres (último recurso)
]


def build_chunker() -> RecursiveCharacterTextSplitter:
    """
    Construye el splitter con parámetros optimizados para RAG.

    chunk_size=512:   Suficiente para contener una idea completa
                      sin exceder el contexto de retrieval típico.
    chunk_overlap=52: ~10% del chunk — preserva contexto en fronteras
                      sin inflar demasiado el índice.

    Trade-off: chunks más grandes (1024+) mejoran comprensión pero
    reducen precisión de recuperación. 512 es el sweet spot empírico
    para la mayoría de pipelines RAG con LLMs modernos.
    """
    return RecursiveCharacterTextSplitter(
        separators=SEMANTIC_SEPARATORS,
        chunk_size=CONFIG.chunk_size,
        chunk_overlap=CONFIG.chunk_overlap,
        length_function=len,
        is_separator_regex=False,
        strip_whitespace=True,
    )


def chunk_blocks(blocks: list[dict], file_path: Path) -> list[dict]:
    """
    Aplica chunking a todos los bloques extraídos de un documento.
    Enriquece cada chunk con metadatos completos para RAG.

    Metadatos indexados en Qdrant:
    - filename, filepath: identificación del documento fuente
    - file_type: pdf / docx
    - file_hash: MD5 para deduplicación
    - chunk_id, total_chunks: posición en el documento
    - page / section: ubicación semántica
    - char_count: tamaño del chunk
    - indexed_at: timestamp ISO
    """
    chunker   = build_chunker()
    file_hash = file_md5(file_path)
    chunks: list[dict] = []
    chunk_idx = 0

    for block in blocks:
        sub_chunks = chunker.split_text(block["text"])

        for sub in sub_chunks:
            sub = sub.strip()
            if len(sub) < CONFIG.min_chunk_length:
                continue

            metadata = {
                "filename":    file_path.name,
                "filepath":    str(file_path),
                "file_type":   file_path.suffix.lower().lstrip("."),
                "file_hash":   file_hash,
                "chunk_id":    chunk_idx,
                "char_count":  len(sub),
                "indexed_at":  datetime.utcnow().isoformat(),
            }

            # Metadatos específicos por tipo
            if "page" in block:
                metadata["page"]        = block["page"]
                metadata["total_pages"] = block.get("total_pages", 1)
            if "section" in block:
                metadata["section"]     = block["section"]
                metadata["section_idx"] = block.get("section_idx", 0)

            chunks.append({"text": sub, "metadata": metadata})
            chunk_idx += 1

    # Actualizar total_chunks en todos los metadatos
    total = len(chunks)
    for ch in chunks:
        ch["metadata"]["total_chunks"] = total

    return chunks


# ══════════════════════════════════════════════════════════════════
# EMBEDDINGS — modelo BGE optimizado para retrieval
# ══════════════════════════════════════════════════════════════════

def load_embedding_model() -> SentenceTransformer:
    """
    Carga el modelo de embeddings.

    BAAI/bge-large-en-v1.5:
    - #1 en MTEB Retrieval benchmark (inglés)
    - 768 dimensiones — balance calidad/tamaño
    - Instrucción especial para indexación vs consulta (ver encode_for_index)

    Alternativa multilingüe: intfloat/multilingual-e5-large
    - 1024 dims, cubre 100+ idiomas incluyendo español
    - Cambiar embedding_dim=1024 en Config

    Trade-off: bge-large-en > multilingual-e5 en inglés puro,
    pero si los documentos son en español, usar multilingual-e5.
    """
    log.info(f"⏳ Cargando modelo de embeddings: {CONFIG.embedding_model}")
    model = SentenceTransformer(CONFIG.embedding_model)
    log.info(f"✅ Modelo cargado ({CONFIG.embedding_dim} dims)")
    return model


# BGE usa instrucciones de representación para retrieval asimétrico:
# - Al INDEXAR: "Represent this passage for searching relevant queries: "
# - Al CONSULTAR: "Represent this query for searching relevant passages: "
BGE_INDEX_PREFIX = "Represent this passage for searching relevant queries: "


def encode_texts(model: SentenceTransformer, texts: list[str], is_query: bool = False) -> list[list[float]]:
    """
    Genera embeddings en batches con el prefijo BGE correcto.

    Retrieval asimétrico BGE:
    - Documentos indexados con prefijo de "passage"
    - Queries con prefijo de "query"
    Esto mejora significativamente la calidad de recuperación.

    normalize_embeddings=True → necesario para cosine similarity con Qdrant.
    """
    if not is_query and "bge" in CONFIG.embedding_model.lower():
        texts = [BGE_INDEX_PREFIX + t for t in texts]

    embeddings = model.encode(
        texts,
        batch_size=CONFIG.batch_size,
        show_progress_bar=False,
        normalize_embeddings=True,   # Obligatorio para cosine similarity
        convert_to_numpy=True,
    )
    return embeddings.tolist()


# ══════════════════════════════════════════════════════════════════
# QDRANT — cliente y gestión de colección
# ══════════════════════════════════════════════════════════════════

def connect_qdrant() -> QdrantClient:
    """Conecta al servidor Qdrant y valida la conexión."""
    log.info(f"🔌 Conectando a Qdrant en {CONFIG.qdrant_host}:{CONFIG.qdrant_port}")
    client = QdrantClient(host=CONFIG.qdrant_host, port=CONFIG.qdrant_port, timeout=30)

    # Validar conexión
    try:
        client.get_collections()
        log.info("✅ Conexión a Qdrant establecida")
    except Exception as e:
        log.error(f"❌ No se pudo conectar a Qdrant: {e}")
        log.error("   Asegúrate de que Docker esté corriendo: docker run -p 6333:6333 qdrant/qdrant")
        raise

    return client


def ensure_collection(client: QdrantClient) -> None:
    """
    Crea la colección si no existe.

    Configuración HNSW optimizada para RAG:
    - Distance.COSINE: funciona con embeddings normalizados
    - m=16: conectividad del grafo (↑m = ↑recall, ↑memoria)
    - ef_construct=200: calidad del índice en build-time
    - on_disk=False: vectores en RAM para máxima velocidad

    Payload indexes: permiten filtrar por metadatos en queries
    (ej: filtrar por filename o file_type antes de vector search)
    """
    existing = [c.name for c in client.get_collections().collections]

    if CONFIG.collection_name in existing:
        log.info(f"📂 Colección '{CONFIG.collection_name}' ya existe")
        return

    log.info(f"🆕 Creando colección '{CONFIG.collection_name}'")

    client.create_collection(
        collection_name=CONFIG.collection_name,
        vectors_config=qdrant_models.VectorParams(
            size=CONFIG.embedding_dim,
            distance=qdrant_models.Distance.COSINE,
            on_disk=False,
        ),
        hnsw_config=qdrant_models.HnswConfigDiff(
            m=CONFIG.hnsw_m,
            ef_construct=CONFIG.hnsw_ef_construct,
            full_scan_threshold=10000,
        ),
        optimizers_config=qdrant_models.OptimizersConfigDiff(
            indexing_threshold=20000,   # Construye HNSW después de N vectores
            memmap_threshold=200000,    # Mueve a mmap si supera este tamaño
        ),
    )

    # Crear índices de payload para filtrado eficiente
    for field_name, schema in [
        ("filename",   qdrant_models.PayloadSchemaType.KEYWORD),
        ("file_type",  qdrant_models.PayloadSchemaType.KEYWORD),
        ("file_hash",  qdrant_models.PayloadSchemaType.KEYWORD),
        ("page",       qdrant_models.PayloadSchemaType.INTEGER),
        ("indexed_at", qdrant_models.PayloadSchemaType.KEYWORD),
    ]:
        client.create_payload_index(
            collection_name=CONFIG.collection_name,
            field_name=field_name,
            field_schema=schema,
        )

    log.info(f"✅ Colección creada con {CONFIG.embedding_dim} dims, HNSW m={CONFIG.hnsw_m}")


def upsert_chunks(
    client: QdrantClient,
    chunks: list[dict],
    embeddings: list[list[float]],
) -> int:
    """
    Sube chunks + embeddings a Qdrant en batches.

    Usa upsert (insert or update) para ser idempotente:
    re-indexar el mismo archivo actualiza los vectores existentes.

    El ID de cada punto se genera como hash determinístico del
    (file_hash + chunk_id) para garantizar idempotencia perfecta.

    Returns: número de puntos indexados
    """
    points: list[qdrant_models.PointStruct] = []

    for chunk, embedding in zip(chunks, embeddings):
        meta = chunk["metadata"]

        # ID determinístico: mismo archivo+chunk siempre mismo ID
        point_id_str = f"{meta['file_hash']}_{meta['chunk_id']}"
        point_id = int(hashlib.md5(point_id_str.encode()).hexdigest(), 16) % (2**63)

        payload = {**meta, "text": chunk["text"]}

        points.append(qdrant_models.PointStruct(
            id=point_id,
            vector=embedding,
            payload=payload,
        ))

    # Subir en batches para no saturar la conexión
    total_upserted = 0
    for i in range(0, len(points), CONFIG.upsert_batch_size):
        batch = points[i : i + CONFIG.upsert_batch_size]
        client.upsert(
            collection_name=CONFIG.collection_name,
            points=batch,
            wait=True,
        )
        total_upserted += len(batch)

    return total_upserted


# ══════════════════════════════════════════════════════════════════
# MANEJO DE ERRORES — mover a carpeta y loguear
# ══════════════════════════════════════════════════════════════════

def handle_error(path: Path, error: Exception) -> None:
    """
    Ante un fallo en un archivo:
    1. Mueve el archivo a la carpeta 'errores' con timestamp
    2. Registra el error detallado en el log
    3. El pipeline principal continúa con el siguiente archivo
    """
    CONFIG.error_folder.mkdir(parents=True, exist_ok=True)

    # Timestamp en nombre para preservar múltiples intentos fallidos
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest_name = f"{path.stem}__{timestamp}{path.suffix}"
    dest_path = CONFIG.error_folder / dest_name

    try:
        shutil.copy2(str(path), str(dest_path))
        log.error(f"❌ Error en '{path.name}': {error}")
        log.error(f"   Archivo copiado a: {dest_path}")
        log.debug(f"   Traceback completo:\n{_get_traceback(error)}")
    except Exception as copy_err:
        log.error(f"❌ Error al copiar '{path.name}' a errores: {copy_err}")


def _get_traceback(error: Exception) -> str:
    """Extrae el traceback formateado de una excepción."""
    import traceback
    return "".join(traceback.format_exception(type(error), error, error.__traceback__))


# ══════════════════════════════════════════════════════════════════
# PIPELINE PRINCIPAL
# ══════════════════════════════════════════════════════════════════

def discover_files(folder: Path) -> list[Path]:
    """
    Descubre todos los archivos PDF y DOCX en la carpeta de entrada,
    incluyendo subcarpetas (rglob).
    """
    supported = {".pdf", ".docx"}
    files = [
        p for p in folder.rglob("*")
        if p.is_file() and p.suffix.lower() in supported
    ]
    files.sort()
    return files


def process_file(
    path: Path,
    model: SentenceTransformer,
    client: QdrantClient,
    state: dict,
) -> Optional[int]:
    """
    Pipeline completo para un archivo:
    1. Verificar si ya fue indexado (hash MD5)
    2. Extraer texto por bloques
    3. Dividir en chunks semánticos
    4. Generar embeddings por batch
    5. Subir a Qdrant

    Returns: número de chunks indexados, o None si fue omitido/error
    """
    file_hash = file_md5(path)

    # Skip si ya fue indexado y no cambió
    if CONFIG.skip_already_indexed and state.get(str(path)) == file_hash:
        log.info(f"  ⏭️  Omitiendo (sin cambios): {path.name}")
        return None

    log.info(f"  📄 Procesando: {path.name}")

    # 1. Extracción de texto
    log.debug(f"     → Extrayendo texto...")
    blocks = extract_text_blocks(path)
    log.debug(f"     → {len(blocks)} bloques extraídos")

    # 2. Chunking
    log.debug(f"     → Aplicando chunking semántico...")
    chunks = chunk_blocks(blocks, path)
    if not chunks:
        raise ValueError("No se generaron chunks válidos tras la extracción")
    log.debug(f"     → {len(chunks)} chunks generados")

    # 3. Embeddings en batches
    log.debug(f"     → Generando embeddings ({len(chunks)} chunks)...")
    texts      = [c["text"] for c in chunks]
    embeddings = encode_texts(model, texts, is_query=False)

    # 4. Upsert en Qdrant
    log.debug(f"     → Indexando en Qdrant...")
    n_indexed = upsert_chunks(client, chunks, embeddings)

    # Actualizar estado persistente
    state[str(path)] = file_hash

    return n_indexed


def run_pipeline() -> None:
    """
    Orquesta el pipeline completo de indexación.
    Gestiona estadísticas, errores y progreso visual.
    """
    start_time = datetime.now()

    log.info("=" * 65)
    log.info("  DOCUMENT INDEXER — Inicio del pipeline")
    log.info(f"  Fecha/Hora : {start_time.strftime('%Y-%m-%d %H:%M:%S')}")
    log.info(f"  Carpeta    : {CONFIG.input_folder.absolute()}")
    log.info(f"  Colección  : {CONFIG.collection_name}")
    log.info(f"  Modelo     : {CONFIG.embedding_model}")
    log.info("=" * 65)

    # Validar carpeta de entrada
    if not CONFIG.input_folder.exists():
        log.error(f"❌ La carpeta '{CONFIG.input_folder}' no existe.")
        log.error("   Crea la carpeta 'Pdf' y añade tus documentos.")
        sys.exit(1)

    # Cargar estado previo
    state = load_state()

    # Descubrir archivos
    files = discover_files(CONFIG.input_folder)
    if not files:
        log.warning(f"⚠️  No se encontraron archivos PDF/DOCX en '{CONFIG.input_folder}'")
        return

    log.info(f"📁 Archivos encontrados: {len(files)}")

    # Inicializar modelo y cliente (una sola vez)
    model  = load_embedding_model()
    client = connect_qdrant()
    ensure_collection(client)

    # Estadísticas
    stats = {
        "total":     len(files),
        "processed": 0,
        "skipped":   0,
        "errors":    0,
        "chunks":    0,
    }

    # Procesar cada archivo con barra de progreso
    log.info("")
    with tqdm(files, desc="Indexando", unit="archivo", ncols=80) as pbar:
        for path in pbar:
            pbar.set_description(f"📄 {path.name[:35]}")
            try:
                result = process_file(path, model, client, state)
                if result is None:
                    stats["skipped"] += 1
                else:
                    stats["processed"] += 1
                    stats["chunks"]    += result
                    log.info(f"  ✅ {path.name} → {result} chunks indexados")
            except Exception as e:
                stats["errors"] += 1
                handle_error(path, e)
                # Continuar con el siguiente archivo (no abortar)
                continue
            finally:
                # Guardar estado después de cada archivo
                save_state(state)

    # Reporte final
    elapsed = (datetime.now() - start_time).total_seconds()
    log.info("")
    log.info("=" * 65)
    log.info("  RESUMEN FINAL")
    log.info(f"  ✅ Procesados   : {stats['processed']}")
    log.info(f"  ⏭️  Omitidos    : {stats['skipped']}  (sin cambios)")
    log.info(f"  ❌ Errores      : {stats['errors']}")
    log.info(f"  🧩 Chunks total : {stats['chunks']}")
    log.info(f"  ⏱️  Tiempo total : {elapsed:.1f}s")

    if stats["errors"] > 0:
        log.info(f"  ⚠️  Archivos con error movidos a: '{CONFIG.error_folder}'")

    # Info de la colección Qdrant
    try:
        info = client.get_collection(CONFIG.collection_name)
        log.info(f"  📊 Puntos en Qdrant: {info.points_count:,}")
    except Exception:
        pass

    log.info("=" * 65)


# ══════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    run_pipeline()
